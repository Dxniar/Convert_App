"""Приложение для конвертации тестов из DOCX в Excel"""

import os
import re
from tkinter import END, Listbox, filedialog, messagebox

import customtkinter as ctk
from docx import Document
from openpyxl import Workbook

ctk.set_appearance_mode("dark")
ctk.set_default_color_theme("blue")


ANSWER_KEYWORDS = (
    "дұрыс жауап",
    "дұрыс жаабы",
    "дұрыс жауаптар",
    "дұрыс жауабы",
    "жауабы",
    "жауап",
    "жауаптар",
    "ответ",
    "ответы",
    "правильный ответ",
    "дұрыс жаурамы",
    "дұрыс жауылы",
)

ANSWER_REGEX = re.compile(r"^[A-Za-zА-Яа-я]\)")


class CleanedDoc:
    def __init__(self, paragraphs):
        self.paragraphs = [
            type('Paragraph', (), {'text': text}) for text in paragraphs]


class DocxToExcelApp(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title("Docx -> Excel конвертер v3.0")
        width = 800
        height = 700

        screen_width = self.winfo_screenwidth()
        screen_height = self.winfo_screenheight()
        x = (screen_width - width) // 2
        y = (screen_height - height) // 2

        self.geometry(f"{width}x{height}+{x}+{y}")
        self.docx_paths = []

        self.label = ctk.CTkLabel(
            self, text="Выберите папку с DOCX файлами:", font=("Arial", 16))
        self.label.pack(pady=10)

        self.btn_browse = ctk.CTkButton(
            self, text="Выбрать папку", command=self.select_folder, font=("Arial", 16))
        self.btn_browse.pack(pady=10)

        self.level_entry_label = ctk.CTkLabel(
            self, text="Уровни сложности (например: A=35,B=50,C=35):", font=("Arial", 14))
        self.level_entry_label.pack(pady=(10, 2))

        self.level_entry = ctk.CTkEntry(self, font=("Arial", 14), width=400)
        self.level_entry.insert(0, "A=35,B=50,C=35")
        self.level_entry.pack(pady=(0, 10))

        self.btn_convert = ctk.CTkButton(
            self, text="Конвертировать в Excel", command=self.convert, state="disabled", font=("Arial", 16))
        self.btn_convert.pack(pady=10)

        self.label_files = ctk.CTkLabel(
            self, text="Файлы в папке с документами:", font=("Arial", 16))
        self.label_files.pack(pady=10)

        frame = ctk.CTkFrame(self)
        frame.pack(padx=20, pady=10, fill="both", expand=True)

        self.listbox = Listbox(
            frame,
            font=("Arial", 14),
            yscrollcommand=lambda *args: self.scrollbar.set(*args),
            bg="#1e1e1e",
            fg="white",
            selectbackground="#3a3a3a",
            selectforeground="white",
            highlightthickness=0,
            bd=0, relief="flat"
        )
        self.listbox.pack(side="left", fill="both", expand=True)
        self.listbox.bind("<Double-Button-1>", self.on_file_double_click)

        self.scrollbar = ctk.CTkScrollbar(
            frame, orientation="vertical", command=self.listbox.yview)
        self.scrollbar.pack(side="right", fill="y")

        self.btn_delete = ctk.CTkButton(
            self, text="Удалить выбранный файл", command=self.delete_selected_file, font=("Arial", 14))
        self.btn_delete.pack(pady=10)

    def select_folder(self):
        folder_path = filedialog.askdirectory()
        if folder_path:
            self.docx_paths = [os.path.join(folder_path, f) for f in os.listdir(
                folder_path) if f.lower().endswith(".docx")]
            self.label.configure(
                text=f"Найдено файлов: {len(self.docx_paths)}")
            self.btn_convert.configure(
                state="normal" if self.docx_paths else "disabled")
            self.show_folder_files(folder_path)

    def show_folder_files(self, folder=None):
        if not folder and self.docx_paths:
            folder = os.path.dirname(self.docx_paths[0])
        if not folder:
            return

        try:
            files = os.listdir(folder)
            files = [f for f in files if f.lower().endswith(('.docx', '.xlsx'))]
            self.listbox.delete(0, END)
            for f in sorted(files):
                self.listbox.insert(END, f)
        except Exception as e:
            self.listbox.insert(END, f"Ошибка при чтении папки: {e}")

    def on_file_double_click(self, event):
        selection = self.listbox.curselection()
        if selection:
            filename = self.listbox.get(selection[0])
            messagebox.showinfo(
                "Инфо", f"Ты просто просматриваешь файл:\n{filename}\n\nДля конвертации выбери папку через кнопку 'Выбрать папку'")

    def delete_selected_file(self):
        selection = self.listbox.curselection()
        if not selection:
            messagebox.showwarning("Внимание", "Файл не выбран")
            return
        filename = self.listbox.get(selection[0])
        for path in self.docx_paths:
            folder = os.path.dirname(path)
            full_path = os.path.join(folder, filename)
            if os.path.exists(full_path):
                answer = messagebox.askyesno(
                    "Удалить файл", f"Вы точно хотите удалить файл?\n{filename}")
                if answer:
                    try:
                        os.remove(full_path)
                        self.docx_paths = [
                            p for p in self.docx_paths if os.path.basename(p) != filename]
                        self.label.configure(text="Файл удалён")
                        self.btn_convert.configure(
                            state="normal" if self.docx_paths else "disabled")
                        self.show_folder_files(folder)
                    except Exception as e:
                        messagebox.showerror(
                            "Ошибка", f"Не удалось удалить файл:\n{e}")
                break

    def starts_with_letter_bracket(self, s: str) -> bool:
        """Проверяет, начинается ли строка с буквы и закрывающей скобки."""
        return bool(ANSWER_REGEX.match(s))

    def start_with_letter_dot(self, line):
        return bool(re.match(r"^[A-Za-z]\.", line))

    def preprocess_lines(self, lines: list[str]) -> list[str]:
        """Разбивает строки вида "вопрос a)" на две отдельные."""
        processed = []
        for line in lines:
            line = line.strip()
            if not line:
                continue
            if not line.lower().startswith(ANSWER_KEYWORDS) and not self.starts_with_letter_bracket(line):
                match = re.search(r"\s([A-Za-zА-Яа-я]\))", line)
                if match:
                    processed.append(line[:match.start(1)].strip())
                    processed.append(line[match.start(1):].strip())
                    continue
            processed.append(line)
        return processed

    def convert(self):
        if not self.docx_paths:
            messagebox.showerror("Ошибка", "Файлы не выбраны")
            return

        try:
            user_input = self.level_entry.get()

            for docx_path in self.docx_paths:
                level_map = {}
                start = 1
                for part in user_input.split(","):
                    level_name, count = part.strip().split("=")
                    num = int(count)
                    end = start + num - 1
                    for i in range(start, end + 1):
                        level_map[i] = level_name
                    start = end + 1

                question_index = 1
                doc_name = os.path.splitext(os.path.basename(docx_path))[0]

                original_doc = Document(docx_path)

                cleaned_paragraphs = []
                for para in original_doc.paragraphs:
                    parts = para.text.split('\n')
                    for part in parts:
                        cleaned_paragraphs.append(part.strip())

                doc = CleanedDoc(cleaned_paragraphs)

                full_text = "\n".join([para.text.strip()
                                       for para in doc.paragraphs if para.text.strip()])
                lines = self.preprocess_lines(full_text.splitlines())

                question_blocks = []
                question_lines = []
                answers = []
                correct_letters = []
                collecting_correct = False

                for line in lines:
                    line = line.strip()
                    if not line:
                        continue
                    line_lower = line.lower()
                    if line_lower.startswith(ANSWER_KEYWORDS):
                        # Проверяем, есть ли варианты в этой же строке после двоеточия
                        parts = line.split(":", 1)
                        if len(parts) > 1 and parts[1].strip():
                            correct_raw = parts[1].strip().lower()
                            # Ищем либо буквы перед скобкой, либо одиночные буквы (например: "c")
                            correct_letters = [m[0].lower() for m in re.findall(
                                r"([A-Za-zА-Яа-яЁё])\)", correct_raw)]
                            if not correct_letters:
                                # Если нет формата "c)", ищем одиночные буквы (например: "c")
                                correct_letters = re.findall(
                                    r"\b([A-Za-zА-Яа-яЁё])\b", correct_raw)
                            collecting_correct = False
                        else:
                            # Следующие строки — правильные варианты
                            correct_letters = []
                            collecting_correct = True
                        continue
                    if collecting_correct:
                        # Пока идут строки вида "a) ..." — собираем буквы
                        m = re.match(r"^([A-Za-zА-Яа-яЁё])\)", line)
                        if m:
                            correct_letters.append(m[1].lower())
                            continue
                        else:
                            collecting_correct = False  # закончились варианты
                    if self.starts_with_letter_bracket(line) or self.start_with_letter_dot(line):
                        answers.append(line)
                    else:
                        if answers:
                            full_question = " ".join(question_lines).strip()
                            question_blocks.append(
                                (full_question, answers.copy(), correct_letters.copy()))
                            question_lines = [line]
                            answers.clear()
                            correct_letters.clear()
                        else:
                            question_lines.append(line)

                if question_lines:
                    full_question = " ".join(question_lines).strip()
                    question_blocks.append(
                        (full_question, answers.copy(), correct_letters.copy()))

                max_answers = max(len(q[1])
                                  for q in question_blocks) if question_blocks else 0
                excel_file = Workbook()
                sheet = excel_file.active
                sheet.title = "Тест"

                header = ["question_text", "level", "is_multiple"]
                for i in range(1, max_answers + 1):
                    header.append(f"answer_{i}")
                    header.append(f"is_correct_{i}")
                sheet.append(header)

                for question, answers, correct_letters in question_blocks:
                    level = level_map.get(question_index, "")
                    is_multiple = 1 if len(correct_letters) > 1 else 0
                    is_correct = []
                    for ans in answers:
                        ans_letter = ans.strip()[0].lower()
                        is_correct.append(
                            1 if ans_letter in correct_letters else 0)
                    row = [question, level, is_multiple]
                    for j in range(max_answers):
                        if j < len(answers) and j < len(is_correct):
                            row.append(answers[j])
                            row.append(is_correct[j])
                        else:
                            row.append("")
                            row.append(0)
                    sheet.append(row)
                    question_index += 1

                xlsx_path = os.path.join(
                    os.path.dirname(docx_path), f"{doc_name}.xlsx")
                excel_file.save(xlsx_path)

            messagebox.showinfo(
                "Готово!", "Все файлы успешно сконвертированы в Excel")
            if self.docx_paths:
                self.show_folder_files(os.path.dirname(self.docx_paths[0]))

        except Exception as e:
            messagebox.showerror("Ошибка", str(e))


if __name__ == "__main__":
    app = DocxToExcelApp()
    app.mainloop()
